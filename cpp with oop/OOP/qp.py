import docx
from docx.shared import Pt, RGBColor

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102)

def add_qa(doc, question, answer, is_numerical=False):
    p = doc.add_paragraph()
    q_run = p.add_run(f"Q: {question}\n")
    q_run.bold = True
    q_run.font.color.rgb = RGBColor(0, 0, 0)
    
    a_run = p.add_run(f"Answer: {answer}")
    if is_numerical:
        a_run.font.color.rgb = RGBColor(153, 0, 0)
        a_run.bold = True
        
        note = doc.add_paragraph()
        note_run = note.add_run("*(Note for numericals: Please watch standard tutorial videos for the detailed step-by-step solution)*")
        note_run.italic = True

def main():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Computer Organization & Architecture - Exam Preparation', 0)
    
    doc.add_paragraph("This document organizes questions from your provided MAKAUT past papers (2023-2026) and aligns them with your syllabus modules. It includes full theoretical answers, final answers for numericals, and predicted suggestions.")

    # Modules 2 & 3
    add_heading(doc, 'Modules 2 & 3: Functional Blocks & Instruction Set Architecture', 1)

    add_qa(doc, "Draw the functional diagram of a computer and explain each block.", 
           "A computer consists of three main functional blocks: CPU, Memory, and I/O.\n- CPU (Central Processing Unit): Contains the ALU (Arithmetic Logic Unit) for calculations and Control Unit (CU) to orchestrate operations.\n- Memory: Stores data and instructions (Primary like RAM for active tasks, Secondary like Disk for storage).\n- I/O Subsystem: Facilitates communication with the outside world (Keyboard, Monitor).\n- System Bus: Connects these components (Data Bus, Address Bus, Control Bus).")

    add_qa(doc, "Explain 0-address, 1-address, 2-address & 3-address instruction formats with suitable examples.",
           "\n- 0-Address: Uses a stack architecture. Operands are implicitly on top of the stack. Example: PUSH A, ADD (pops two, adds, pushes result).\n- 1-Address: Uses an Accumulator (AC) register. One operand is in AC, the other in memory. Example: ADD B (AC <- AC + M[B]).\n- 2-Address: Two operands are specified; destination and source. Example: ADD R1, R2 (R1 <- R1 + R2).\n- 3-Address: Three operands specified (two sources, one destination). Example: ADD R1, R2, R3 (R1 <- R2 + R3).")

    add_qa(doc, "What are the different types of addressing modes in computer architecture? Provide examples.",
           "\n1. Immediate: Operand is in the instruction (e.g., ADD R1, #5).\n2. Direct: Address of operand is in the instruction (e.g., LOAD R1, 1000).\n3. Indirect: Address field points to a memory location containing the effective address.\n4. Register: Operand is in a register (e.g., ADD R1, R2).\n5. Register Indirect: Register contains the address of the operand.\n6. Displacement/Indexed: Effective address is base register + offset.")

    add_qa(doc, "What is the difference between RISC and CISC architectures?",
           "\n- RISC (Reduced Instruction Set Computer): Few, simple instructions, fixed-length format, executes in one clock cycle, highly pipelined, many registers, uses Load/Store architecture.\n- CISC (Complex Instruction Set Computer): Many complex instructions, variable-length formats, takes multiple cycles per instruction, fewer registers, instructions can directly manipulate memory.")

    # Modules 4 & 5
    add_heading(doc, 'Modules 4 & 5: Data Representation & Computer Arithmetic', 1)

    add_qa(doc, "Explain IEEE 754 standard format for floating point representation.",
           "IEEE 754 defines formats for floating-point numbers.\n- Single Precision (32 bits): 1 sign bit, 8 bits for biased exponent (bias 127), 23 bits for normalized mantissa (fraction).\n- Double Precision (64 bits): 1 sign bit, 11 bits for biased exponent (bias 1023), 52 bits for normalized mantissa.\nThe value is calculated as: (-1)^S * (1.Fraction) * 2^(Exponent - Bias).")

    add_qa(doc, "What are the advantages of Carry Look Ahead (CLA) over ripple carry adder?",
           "In a Ripple Carry Adder, the carry bit ripples from the least significant bit to the most significant bit, causing a propagation delay proportional to the number of bits (O(n)). \nA Carry Look Ahead (CLA) adder reduces this delay by calculating the carry bits in advance using 'Generate' (G = A AND B) and 'Propagate' (P = A XOR B) functions. Its delay is independent of the number of bits (O(1) logic depth), making it significantly faster for large bit-width additions.")

    add_qa(doc, "Multiply -9 by +6 using Booth's algorithm. Assume 5 bits each. (Numerical)",
           "Final Answer: Binary 2's complement result is 111001010 (-54 in decimal).", True)

    # Module 7
    add_heading(doc, 'Module 7: CPU Control Unit Design', 1)

    add_qa(doc, "Differentiate between hardwired control and microprogrammed control.",
           "\n- Hardwired Control: Implemented using combinational logic gates and flip-flops. It is very fast but inflexible. Any change requires rewiring hardware. Used in RISC processors.\n- Microprogrammed Control: Control signals are generated by reading microinstructions from a Control Memory (ROM). It is slower than hardwired but highly flexible; modifications just require changing the software (microcode). Used in CISC processors.")

    # Module 9
    add_heading(doc, 'Module 9: Peripheral Devices', 1)

    add_qa(doc, "Explain the role of DMA (Direct Memory Access) and the DMA Controller.",
           "DMA allows peripheral devices to transfer data directly to or from main memory without the continuous intervention of the CPU. The CPU initiates the transfer by giving the DMA controller the start address, word count, and direction. The CPU then does other work. The DMA controller manages the buses and transfers the data, sending an interrupt to the CPU only when the entire block transfer is complete. This significantly improves system throughput.")

    add_qa(doc, "Explain daisy chaining priority interrupts vs Polling.",
           "\n- Daisy Chaining: Hardware-based priority. Devices are connected in a serial chain. The interrupt acknowledge signal propagates through the chain; the first requesting device intercepts it. Fast, but hard to modify priorities. If a device fails, the chain breaks.\n- Polling: Software-based. The CPU executes a routine checking each device's status register in a specific order to find which generated the interrupt. Flexible, but wastes CPU cycles checking devices.")

    # Module 10
    add_heading(doc, 'Module 10: Pipelining & Parallel Processors', 1)

    add_qa(doc, "Explain the concept of pipeline hazards. How can they be avoided?",
           "Pipeline hazards prevent the next instruction from executing during its designated clock cycle.\n1. Structural Hazards: Resource conflicts (e.g., ALU needed by two stages). Avoided by duplicating hardware (separate data/instruction caches).\n2. Data Hazards: An instruction depends on the result of a previous incomplete instruction (RAW, WAR, WAW). Avoided by Data Forwarding (bypassing) and pipeline stalling (bubbles).\n3. Control Hazards: Caused by branch instructions changing the PC. Avoided by Branch Prediction, delayed branching, and flushing.")

    add_qa(doc, "What is cache coherence? Explain the MESI protocol briefly.",
           "In multiprocessor systems where each CPU has its own cache, cache coherence ensures that if one CPU updates a shared memory variable, other CPUs see the updated value, preventing inconsistent data views.\nMESI is a snooping protocol with 4 states for cache lines:\n- Modified (M): Line is modified and only in this cache.\n- Exclusive (E): Line is clean and only in this cache.\n- Shared (S): Line is clean and present in multiple caches.\n- Invalid (I): Line contains invalid data.")

    add_qa(doc, "Given a 4 segment pipeline (40ns, 25ns, 45ns, 45ns) + 5ns interface register. Calculate cycle time, execution time for 100 tasks, real speedup, max speedup. (Numerical)",
           "Cycle time (tp) = Max(40, 25, 45, 45) + 5 = 50 ns.\nNon-pipeline time (tn) = 40+25+45+45 = 155 ns.\nExecution time (100 tasks) = (k + n - 1) * tp = (4 + 100 - 1) * 50 = 103 * 50 = 5150 ns.\nReal Speedup = (n * tn) / ((k+n-1)*tp) = (100 * 155) / 5150 = 3.009.\nMax Speedup = k = 4.", True)

    # Module 11
    add_heading(doc, 'Module 11: Memory Organization', 1)

    add_qa(doc, "Explain the concept of virtual memory and its advantages.",
           "Virtual memory is a technique that gives the illusion of a main memory much larger than the physical RAM available. It uses secondary storage (like a hard disk) to store parts of running programs, loading pages into RAM only when needed (demand paging). \nAdvantages: Allows execution of programs larger than physical memory, enables efficient multiprogramming, and provides memory isolation/protection between processes.")

    add_qa(doc, "Compare Direct Mapped, Set Associative, and Fully Associative cache mapping.",
           "\n- Direct Mapped: Each main memory block maps to exactly one cache line (Line = Block Number MOD Number of Lines). Simple, fast, but suffers from high conflict misses.\n- Fully Associative: A memory block can be placed in ANY cache line. Requires searching all tags simultaneously. Very flexible, lowest miss rate, but hardware is complex and expensive.\n- Set Associative: Cache is divided into sets. A block maps to a specific set, but can be placed anywhere within that set. A compromise providing good hit rates with reasonable hardware complexity.")

    add_qa(doc, "What is page fault? Explain replacement algorithms (LRU, FIFO, Optimal).",
           "A page fault occurs when the CPU attempts to access a page that is part of the virtual address space but is not currently loaded in physical RAM. The OS must fetch it from the disk.\n- FIFO (First In First Out): Evicts the oldest page. Simple but suffers from Belady's Anomaly (more frames = more faults).\n- LRU (Least Recently Used): Evicts the page that has not been accessed for the longest time. Good performance, realistic.\n- Optimal: Evicts the page that will not be used for the longest time in the future. Impossible to implement (requires future knowledge) but used as a benchmark.")

    add_qa(doc, "Calculate TAG, LINE, and WORD bits for a 4-way set associative cache (128 lines, 64 words/line) with 20-bit address. (Numerical)",
           "WORD = 6 bits (since 2^6 = 64 words).\nSets = Total lines / ways = 128 / 4 = 32 sets.\nSET/LINE = 5 bits (since 2^5 = 32).\nTAG = Total Address - (Set + Word) = 20 - (5 + 6) = 9 bits.\nFinal Answer: TAG=9, SET=5, WORD=6.", True)

    # Gemini's Suggestions
    add_heading(doc, "Gemini's Top Suggestions (Highly Probable Topics for Tomorrow)", 1)
    
    add_qa(doc, "Explain Restoring vs Non-Restoring Division Algorithm.", 
           "Review the flowcharts for both. Restoring adds the divisor back if the partial remainder is negative; non-restoring skips the restore step and changes the operation (add/sub) in the next step. Expect a numerical to trace the steps.")
    add_qa(doc, "Amdahl's Law and Speedup calculations", 
           "Speedup = 1 / ((1 - F) + F/N), where F is the parallelizable fraction and N is the number of processors. Very common short numerical.")
    add_qa(doc, "Cache Write Policies (Write-through vs Write-back)", 
           "Write-through updates RAM and Cache simultaneously (slow but safe). Write-back updates Cache only, sets a 'dirty bit', and updates RAM only when the block is evicted (fast, requires coherence protocols).")

    # Save the document
    filename = 'COA_Exam_Prep_MAKAUT.docx'
    doc.save(filename)
    print(f"Document saved successfully as {filename}")

if __name__ == "__main__":
    main()